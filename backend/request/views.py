import csv
import itertools
import json
import logging
import os
from io import StringIO
from unicodedata import normalize
from urllib.parse import urlencode

from common.serializers import UserSerializer
from common.utils import retrieve_group_items
from common.views import CsrfExemptSessionAuthentication, StandardResultsSetPagination
from django.apps import apps
from django.conf import settings
from django.contrib.auth import get_user_model
from django.contrib.auth.decorators import login_required
from django.contrib.sites.shortcuts import get_current_site
from django.core.exceptions import PermissionDenied
from django.core.mail import send_mail
from django.db import transaction
from django.db.models import Prefetch
from django.http import (
    Http404,
    HttpRequest,
    HttpResponse,
    HttpResponseRedirect,
    JsonResponse,
)
from django.shortcuts import get_object_or_404, render
from django.template.loader import render_to_string
from django.urls import reverse
from django.utils import dateformat, timezone
from django.utils.crypto import get_random_string
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt
from fpdf import FPDF, HTMLMixin
from fpdf.errors import FPDFUnicodeEncodingException
from library_sample_shared.models import LibraryProtocol
from library_sample_shared.serializers import LibraryProtocolSerializer
from rest_framework import filters, status, viewsets
from rest_framework.decorators import action
from rest_framework.exceptions import NotFound
from rest_framework.permissions import AllowAny, IsAdminUser
from rest_framework.response import Response
from tablib import Dataset

from .models import FileRequest, Request
from .resources import LibrariesResource, SamplesResource
from .serializers import RequestFileSerializer, RequestSerializer

User = get_user_model()
Library = apps.get_model("library", "Library")
Sample = apps.get_model("sample", "Sample")
LibraryPreparation = apps.get_model("library_preparation", "LibraryPreparation")

logger = logging.getLogger("db")


class PDF(FPDF):  # pragma: no cover
    def __init__(self, title="Title", font="Arial"):
        self.core_fonts_encoding = "UTF-8"
        self.title = title
        self.font = font
        super().__init__()

    def header(self):
        self.set_font(self.font, style="B", size=14)  # Arial bold 15
        self.cell(0, 10, self.title, align="C")  # Title
        self.ln(10)  # Line break

    def footer(self):
        self.set_y(-15)  # Position at 1.5 cm from bottom
        self.set_font(self.font, size=8)  # Arial 8
        # Page number
        self.cell(0, 10, "Page " + str(self.page_no()) + " of {nb}", 0, 0, "C")

    def info_row(self, title, value):
        self.set_font(self.font, style="B", size=11)
        self.cell(35, 10, title + ":")
        self.set_font(self.font, size=11)
        self.cell(0, 10, value)
        self.ln(6)

    def multi_info_row(self, title, value):
        self.set_font(self.font, style="B", size=11)
        self.ln(3)
        self.cell(35, 4, title + ":")
        self.set_font(self.font, size=11)
        self.multi_cell(0, 5, value)
        self.ln(6)

    def multi_checkbox_row(self, title, values):
        self.set_font(self.font, style="B", size=11)
        self.ln(3)
        self.cell(35, 4, title + ":")
        for i in range(len(values)):
            if i > 0:
                self.cell(35, 4, "")
            self.set_font("glyphicons", size=11)
            self.cell(1, 4, "")
            self.set_font(self.font, size=11)
            self.multi_cell(0, 5, values[i])
            self.ln(1)
        self.ln(1)

    def table_row(self, index, name, barcode, type, depth, bold=False):
        if bold:
            self.set_font(self.font, style="B", size=11)
        else:
            self.set_font(self.font, size=11)
        self.cell(10, 10, str(index))
        self.cell(60, 10, name)
        self.cell(40, 10, barcode)
        self.cell(35, 10, type)
        self.cell(0, 10, str(depth))
        self.ln(6)


class Report(FPDF, HTMLMixin):
    def __init__(self, title="Report", font="Arial"):
        self.core_fonts_encoding = "UTF-8"
        self.title = title
        self.font = font
        super().__init__()

    def header(self):
        self.set_font(family=self.font, size=8)
        self.set_text_color(r=189, g=189, b=189)
        self.cell(0, 10, "COMPLETE REPORT", align="L")
        self.cell(0, 10, "Deep Sequencing Facility @ MPI-IE, Freiburg", align="R")
        self.ln(10)

    def footer(self):
        self.set_y(-15)  # Position at 1.5 cm from bottom
        self.set_font(self.font, size=8)  # Arial 8
        # Page number
        self.cell(0, 10, "Page " + str(self.page_no()) + " of {nb}", 0, 0, "C")

    def page_header(self, text):
        self.set_font(family=self.font, style="B", size=12)
        self.cell(0, 10, text)
        self.ln(14)

    def text_block(self, text, style="", size=11, multi=False):
        self.set_font(family=self.font, style=style, size=size)
        if multi:
            self.multi_cell(0, 6, text)
        else:
            self.cell(0, 10, text)
        self.ln(6)

    def generate_html_table(self, data):
        if len(data) == 0:
            return ""

        columns = list(data[0].keys())
        length = len(columns)

        thead = "".join(
            map(
                lambda c: f'<th width="{100 // length}%" align="left">{c}</th>',
                columns,
            )
        )

        tbody = []
        for item in data:
            row = "".join(map(lambda x: f"<td>{x}</td>", item.values()))
            tbody.append(f"<tr>{row}</tr>")
        tbody = "".join(tbody)

        html = """
        <font face="Arial" size="10">
             <table border="0" width="100%">
                <thead>
                    <tr>{}</tr>
                </thead>
                <tbody>{}</tbody>
            </table>
        </font>
        """.format(
            thead, tbody
        )
        html = html.replace("\n", "")

        return html


class RequestViewSet(viewsets.ModelViewSet):
    serializer_class = RequestSerializer
    pagination_class = StandardResultsSetPagination

    filter_backends = (filters.SearchFilter,)
    search_fields = (
        "name",
        "description",
        "user__first_name",
        "user__last_name",
        "user__pi__name",
        "user__organization__name",
        "cost_unit__name",
    )

    def get_queryset(self, showAll=True):
        libraries_qs = Library.objects.all().only("status", "sequencing_depth")
        samples_qs = Sample.objects.all().only("status", "sequencing_depth")
        #   print(libraries_qs.values())

        queryset = (
            Request.objects.filter(archived=False)
            .select_related("user")
            .prefetch_related(
                Prefetch("libraries", queryset=libraries_qs),
                Prefetch("samples", queryset=samples_qs),
                "files",
            )
            .order_by("-create_time")
        )

        if not showAll:
            queryset = queryset.filter(sequenced=False)
        if self.request.user.is_staff:
            pass
        else:
            if not self.request.user.is_pi:
                queryset = queryset.filter(user=self.request.user)
            else:
                queryset = retrieve_group_items(self.request, queryset)

        return queryset

    def list(self, request):
        """Get the list of requests."""

        showAll = False
        if request.GET.get("showAll") == "True":
            showAll = True

        queryset = self.filter_queryset(self.get_queryset(showAll))

        try:
            page = self.paginate_queryset(queryset)
        except NotFound:
            page = None

        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)

        return Response(serializer.data)

    def create(self, request):
        """Create a request."""
        post_data = self._get_post_data(request)
        post_data.update({"user": request.user.pk})
        serializer = self.serializer_class(data=post_data)

        if serializer.is_valid():
            serializer.save()
            return Response({"success": True, "pk": serializer.data["pk"]}, 201)

        else:
            return Response(
                {
                    "success": False,
                    "message": "Invalid payload.",
                    "errors": serializer.errors,
                },
                400,
            )

    @action(methods=["post"], detail=True)
    def edit(self, request, pk=None):
        """Update request with a given id."""
        instance = self.get_object()
        post_data = self._get_post_data(request)

        post_data.update({"user": instance.user.pk})

        serializer = self.get_serializer(data=post_data, instance=instance)

        if serializer.is_valid():
            serializer.save()
            return Response({"success": True})

        else:
            return Response(
                {
                    "success": False,
                    "message": "Invalid payload.",
                    "errors": serializer.errors,
                },
                400,
            )

    @action(methods=["post"], detail=True)
    def mark_as_complete(self, request, pk=None):
        """Mark request as complete, set sequenced to true"""

        instance = Request.objects.filter(archived=False, pk=pk)

        post_data = self._get_post_data(request)
        override = post_data["override"]

        if post_data["override"] == "False":
            override = False
        else:
            override = True

        def checkifcomplete(element):
            if element == 5:
                return True
            else:
                return False

        if override:
            # print("Override is true")
            instance.update(sequenced=True)
            return Response({"success": True})

        else:
            # print("Override is false")
            # print(instance.statuses)
            # check if all libraries/samples related to this requested have been sequenced
            statuses = [status for x in instance for status in x.statuses]

            complete = all([checkifcomplete(x) for x in statuses])

            if complete:
                # print("all statuses are complete")
                instance.update(sequenced=True)
                return Response({"success": True})
            elif not complete:
                # print("there are incomplete statuses")
                return Response({"noncomplete": True})
            else:
                return Response({"error": "error"})

    @action(methods=["post"], detail=True)
    def samples_submitted(self, request, pk=None):
        instance = self.get_object()
        post_data = self._get_post_data(request)
        instance.samples_submitted = post_data["result"]
        instance.save(update_fields=["samples_submitted"])
        return Response({"success": True})

    @action(methods=["get"], detail=True)
    def get_records(self, request, pk=None):
        """Get the list of record's submitted libraries and samples."""
        libraries_qs = Library.objects.all().only(
            "name",
            "barcode",
        )
        samples_qs = Sample.objects.all().only(
            "name",
            "barcode",
            "is_converted",
        )

        try:
            instance = (
                Request.objects.filter(archived=False, pk=pk)
                .prefetch_related(
                    Prefetch("libraries", queryset=libraries_qs),
                    Prefetch("samples", queryset=samples_qs),
                )
                .only("libraries", "samples")
                .first()
            )

            data = [
                {
                    "pk": obj.pk,
                    "name": obj.name,
                    "barcode": obj.barcode,
                    "record_type": obj.__class__.__name__,
                    "is_converted": (
                        True
                        if hasattr(obj, "is_converted") and obj.is_converted
                        else False
                    ),
                }
                for obj in instance.records
            ]

            data = sorted(data, key=lambda x: x["barcode"][3:])
            Res = Response(data)

        except AttributeError as e:
            Res = Response({"success": False, "message": f"{pk} not found!"}, 400)
        return Res

    @action(methods=["get"], detail=True)
    def get_protocol(self, request, pk=None):
        """For example, to poll IDs of Nanopore's Sequencing Kits"""
        instance = LibraryProtocol(id=pk)
        serializer = LibraryProtocolSerializer(instance)
        return Response(serializer.data)

    @action(methods=["get"], detail=True)
    def get_contact_details(self, request, pk=None):
        """Get the user contact details."""
        users_qs = User.objects.all()
        data = (
            Request.objects.filter(pk=pk)
            .prefetch_related(Prefetch("user", queryset=users_qs))
            .only("user")
            .first()
        )
        serializer = UserSerializer(data.user)
        return Response(serializer.data)

    @action(methods=["get"], detail=True)
    def get_files(self, request, pk=None):
        """Get the list of attached files for a request with a given id."""
        instance = self.get_object()
        files = instance.files.all().order_by("name")
        serializer = RequestFileSerializer(files, many=True)
        return Response(serializer.data)

    @action(
        methods=["post"],
        detail=False,
        authentication_classes=[CsrfExemptSessionAuthentication],
    )
    def upload_files(self, request):
        file_ids = []

        if not any(request.FILES):
            return JsonResponse(
                {"success": False, "message": "No files provided."}, status=400
            )

        for file in request.FILES.getlist("files"):
            f = FileRequest(name=file.name, file=file)
            f.save()
            file_ids.append(f.id)

        return JsonResponse({"success": True, "fileIds": file_ids})

    @action(methods=["get"], detail=False)
    def get_files_after_upload(self, request):
        file_ids = json.loads(request.query_params.get("file_ids", "[]"))
        error = ""
        data = []

        try:
            files = [f for f in FileRequest.objects.all() if f.id in file_ids]

            data = [
                {
                    "id": file.id,
                    "name": file.name,
                    "size": file.file.size,
                    "path": settings.MEDIA_URL + file.file.name,
                }
                for file in files
            ]

        except Exception as e:
            error = "Could not get the attached files."
            logger.exception(e)

        return JsonResponse(
            {
                "success": not error,
                "error": error,
                "data": data,
            }
        )

    @action(methods=["get"], detail=False)
    def download_RELACS_Pellets_Abs_form(self, request):
        file_path = os.path.join(settings.STATIC_ROOT, "docs/RELACS.xlsx")

        with open(file_path, "rb") as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-excel")
            response["Content-Disposition"] = "inline; filename=" + os.path.basename(
                file_path
            )
            return response

    @action(methods=["get"], detail=True)
    def download_deep_sequencing_request(self, request, pk=None):  # pragma: no cover
        """Generate a deep sequencing request form in PDF."""
        instance = self.get_object()
        user = instance.user
        organization = user.organization.name if user.organization else ""
        cost_unit = instance.cost_unit.name if instance.cost_unit else ""
        objects = list(
            itertools.chain(
                instance.samples.all(),
                instance.libraries.all(),
            )
        )
        records = [
            {
                "name": obj.name,
                "type": obj.__class__.__name__,
                "barcode": obj.barcode,
                "depth": obj.sequencing_depth,
            }
            for obj in objects
        ]
        records = sorted(records, key=lambda x: x["barcode"][3:])

        declaration_general = "None of the samples listed below are potentially infectious. The listed samples can be handled in an S1 laboratory (BioSafety Level 1) without any safety concerns."

        declaration_gmo = [
            (
                "  Non-GMO samples. Samples listed below do not fall under GenTG regulation (naked DNA, Sequencing libraries, RNA, proteins or metabolites, fixed or lysed cells, etc.). No additional documentation is required."
            ),
            (
                '  GMO samples of BioSafety Level 1 (BSL 1). Samples listed below fall under GenTG regulation. Additional documentation is required: an electronic copy of "Formblatt S1" (in "white" folder) in editable format is provided to the Deep Sequencing Facility BEFORE bringing S1 GMO to the facility (e.g. upload "Formblatt S1" via Parkour to your Request).'
            ),
            (
                "  GMO samples of BioSafety Level 2 (BSL 2). Not possible to be processed in the Deep Sequencing Facility."
            ),
        ]

        pdf = PDF("Deep Sequencing Request")
        pdf.add_font(
            "glyphicons",
            "",
            "/usr/src/app/static/fonts/glyphicons-halflings-regular.ttf",
            uni=True,
        )
        pdf.set_draw_color(217, 217, 217)
        pdf.alias_nb_pages()
        pdf.add_page()

        # Deep Sequencing Request info
        pdf.info_row("Request Name", instance.name)
        pdf.info_row("Date", instance.create_time.strftime("%d.%m.%Y"))
        pdf.info_row("User", user.full_name)
        pdf.info_row("Phone", user.phone if user.phone else "")
        pdf.info_row("Email", user.email)
        pdf.info_row("Organization", organization)
        pdf.info_row("Cost Unit", cost_unit)
        pdf.multi_info_row("Declaration", declaration_general)
        pdf.multi_checkbox_row("GMO Samples", declaration_gmo)
        try:
            pdf.multi_info_row("Description", instance.description)
        except FPDFUnicodeEncodingException:
            pdf.multi_info_row("Description", "ERROR: Character-set outside UTF-8.")
        # except:
        # pdf.multi_info_row(f"Description", "ERROR: {Exception}")

        y = pdf.get_y()
        pdf.line(pdf.l_margin + 1, y, pdf.w - pdf.r_margin - 1, y)

        # List of libraries/samples
        heading = "List of libraries/samples to be submitted for sequencing"
        pdf.ln(5)
        pdf.cell(0, 10, heading, align="C")
        pdf.ln(10)

        pdf.table_row("#", "Name", "Barcode", "Type", "Sequencing Depth (M)", True)

        for i, record in enumerate(records):
            pdf.table_row(
                i + 1,
                record["name"],
                record["barcode"],
                record["type"],
                record["depth"],
            )

        pdf.ln(10)
        y = pdf.get_y()
        pdf.line(pdf.l_margin + 1, y, pdf.w - pdf.r_margin - 1, y)
        pdf.ln(30)

        # Ensure there is enough space for the signature
        if pdf.get_y() > 265:
            pdf.add_page()
            pdf.ln(20)

        # Signature
        pdf.set_draw_color(0, 0, 0)
        y = pdf.get_y()
        x1_date = pdf.w / 2
        x2_date = x1_date + 45
        x1_signature = x2_date + 5
        x2_signature = pdf.w - pdf.r_margin - 1
        pdf.line(x1_date, y, x2_date, y)
        pdf.line(x1_signature, y, x2_signature, y)

        pdf.set_x(x1_date + (x2_date - x1_date) / 2 - 6)
        pdf.cell(12, 10, "(Date)")
        pdf.set_x(x1_signature + 2)
        pdf.cell(0, 10, "(Principal Investigator)")

        # Generate response
        request_name = (
            normalize("NFKD", instance.name).encode("ASCII", "ignore").decode("utf-8")
        )
        f_name = request_name + "_Deep_Sequencing_Request.pdf"
        response = HttpResponse(bytes(pdf.output()), content_type="application/pdf")
        response["Content-Disposition"] = 'attachment; filename="%s"' % f_name

        return response

    @action(
        methods=["post"],
        detail=True,
        authentication_classes=[CsrfExemptSessionAuthentication],
    )
    def upload_deep_sequencing_request(self, request, pk=None):
        """
        Upload a deep sequencing request with the PI's signature and
        change request's libraries' and samples' statuses to 1.
        """
        instance = self.get_object()

        if not any(request.FILES):
            return JsonResponse(
                {"success": False, "message": "File is missing."}, status=400
            )

        instance.deep_seq_request = request.FILES.get("file")
        instance.save()

        file_name = instance.deep_seq_request.name.split("/")[-1]
        file_path = settings.MEDIA_URL + instance.deep_seq_request.name

        instance.libraries.all().update(status=1)
        instance.samples.all().update(status=1)

        return JsonResponse({"success": True, "name": file_name, "path": file_path})

    @action(methods=["post"], detail=True)
    def solicit_approval(self, request, pk=None):  # pragma: no cover
        """Send an email to the PI."""
        error = ""
        instance = self.get_object()
        subject = f"[ Parkour2 | sequencing experiment is pending approval ] "
        subject += request.data.get("subject", "")
        message = request.data.get("message", "")
        include_records = json.loads(request.POST.get("include_records", "true"))
        records = []
        try:
            if instance.user.pi.archived:
                raise ValueError(
                    "PI: "
                    + instance.user.pi.name
                    + ", is no longer enrolled. Please contact an admin."
                )
            elif instance.user.pi.email == "Unset":
                raise ValueError(
                    "PI: "
                    + instance.user.pi.name
                    + ", has no e-mail address assigned. Please contact an admin."
                )
            records = list(instance.libraries.all()) + list(instance.samples.all())
            for r in records:
                if r.status != 0:
                    raise ValueError("Not all records have status of zero.")
            records = sorted(records, key=lambda x: x.barcode[3:])
            if not include_records:
                records = []
            instance.token = get_random_string(30)
            instance.save(update_fields=["token"])
            url_scheme = request.is_secure() and "https" or "http"
            url_domain = get_current_site(request).domain
            url_query = urlencode({"token": instance.token})
            send_mail(
                subject=subject,
                message="",
                html_message=render_to_string(
                    "approval.html",
                    {
                        "full_name": instance.user.full_name,
                        "pi_name": instance.user.pi.name,
                        "message": message,
                        "token_url": f"{url_scheme}://{url_domain}/api/approve/this/?pk={instance.id}&{url_query}",
                        "records": records,
                    },
                ),
                from_email=settings.SERVER_EMAIL,
                recipient_list=[instance.user.pi.email],
            )
        except Exception as e:
            error = str(e)
            logger.exception(e)
        return JsonResponse({"success": not error, "error": error})

    @action(methods=["post"], detail=True, permission_classes=[IsAdminUser])
    def send_email(self, request, pk=None):  # pragma: no cover
        """Send an email to the user."""
        error = ""

        instance = self.get_object()
        subject = request.data.get("subject", "")
        message = request.data.get("message", "")
        include_failed_records = json.loads(
            request.POST.get("include_failed_records", "false")
        )
        records = []

        try:
            if subject == "" or message == "":
                raise ValueError("Email subject and/or message is missing.")

            if include_failed_records:
                records = list(instance.libraries.filter(status=-1)) + list(
                    instance.samples.filter(status=-1)
                )
                records = sorted(records, key=lambda x: x.barcode[3:])

            send_mail(
                subject=subject,
                message="",
                html_message=render_to_string(
                    "email.html",
                    {
                        "full_name": instance.user.full_name,
                        "message": message,
                        "records": records,
                    },
                ),
                from_email=settings.SERVER_EMAIL,
                recipient_list=[instance.user.email],
            )

        except Exception as e:
            error = str(e)
            logger.exception(e)

        return JsonResponse({"success": not error, "error": error})

    @action(methods=["get"], detail=True)
    def download_complete_report(self, request, pk=None):
        def add_table(document, header, data, contains_comments=True):
            if contains_comments:
                # Create table
                # table = document.add_table(rows=1, cols=len(header))
                table = document.add_table(rows=1, cols=len(header) - 1)
                hdr_cells = table.rows[0].cells
                # for i, h in enumerate(header):
                for i, h in enumerate(header[:-1]):
                    hdr_cells[i].text = h
                for row in data:
                    row_cells = table.add_row().cells
                    # for i, value in enumerate(row):
                    for i, value in enumerate(row[:-1]):
                        row_cells[i].text = str(value)
                    # add comment row
                    comment_cells = table.add_row().cells
                    comment_cells[0].merge(comment_cells[-1])
                    comment_cells[0].text = "Comments: " + str(row[-1])

                # Change font size for all cells
                for row in table.rows:
                    row.height = Cm(0.7)
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size = Pt(9)
            else:
                # Create table
                table = document.add_table(rows=1, cols=len(header))
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(header):
                    hdr_cells[i].text = h
                for row in data:
                    row_cells = table.add_row().cells
                    # for i, value in enumerate(row):
                    for i, value in enumerate(row[:-1]):
                        row_cells[i].text = str(value)

                # Change font size for all cells
                for row in table.rows:
                    row.height = Cm(0.7)
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size = Pt(9)

        f_name = "QC Complete Report.docx"
        response = HttpResponse(
            content_type="application/vnd.openxmlformats"
            + "-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="{f_name}"'

        instance = self.get_object()
        records = sorted(
            list(
                itertools.chain(
                    instance.libraries.all(),
                    instance.samples.all(),
                )
            ),
            key=lambda x: x.barcode[3:],
        )

        # Create DOCX document and set default font family
        doc = Document()
        font = doc.styles["Normal"].font
        font.name = "Arial"

        # Page 1
        doc.add_heading("Complete Report", 0)
        p = doc.add_paragraph("")
        p.add_run("Date, Request ID").bold = True

        doc.add_paragraph("")

        p = doc.add_paragraph("")
        p.add_run("Table of Contents").bold = True
        doc.add_paragraph("Summary")
        doc.add_paragraph("Quality Control of received samples")
        doc.add_paragraph("Library Construction")
        doc.add_paragraph("Cluster Generation and Sequencing")
        doc.add_paragraph("Acknowledgements")
        doc.add_paragraph("Appendix")
        doc.add_page_break()

        # Page 2
        doc.add_heading("General Summary of Workflow", 1)
        doc.add_paragraph()
        doc.add_paragraph(
            "Submitted samples or libraries undergo an "
            + "incoming quality control using appropriate "
            + "analytical instruments (Fluorometer, Capillary "
            + "Electrophoresis, qPCR etc). All samples that "
            + "pass international quality standards are "
            + "subjected to appropriate library preparation "
            + "methods. Qualified libraries are pooled for "
            + "multiplex sequencing. An Index Generator "
            + "Software assures suitable index design. Pooled "
            + "libraries are sequenced to reach desired "
            + "depth/coverage using installed sequencing "
            + "instruments. Immediately after the sequencing "
            + "run bcl to fastq conversion and demultiplexing "
            + "is done and the user informed."
        )
        doc.add_page_break()

        # Page 3
        doc.add_heading("Quality Control of received samples/libraries", 1)
        doc.add_paragraph()
        doc.add_paragraph(
            "All documented measurements were conducted by "
            + "the deep sequencing facility, MPI-IE Freiburg. "
            + "Raw data and reports of fluorometric "
            + "quantification (Qubit) and size distribution "
            + "measurements (Fragment Analyzer) can be found as "
            + "attachment to each request in Parkour "
            + "(parkour.ie-freiburg.mpg.de)."
        )
        header = [
            "Date",
            "ID",
            "Name",
            "L/S",
            "Nuc.Type",
            "ng/µl",
            "bp",
            "Comments",
        ]
        data = []
        for r in records:
            rtype = r.__class__.__name__
            row = [
                r.create_time.strftime("%d.%m.%Y"),
                r.barcode,
                r.name,
                rtype[0],
                r.nucleic_acid_type.name if rtype == "Sample" else "-",
                r.concentration,
                r.mean_fragment_size if rtype == "Library" else "-",
                r.comments,
            ]
            data.append(row)
        add_table(doc, header, data)
        doc.add_page_break()

        # Page 4
        doc.add_heading("Library Construction", 1)
        doc.add_paragraph()
        doc.add_paragraph(
            "Documentation is only possible if libraries were "
            + "constructed in the deep sequencing facility, "
            + "MPI-IE Freiburg. Raw data and reports of "
            + "fluorometric quantification (Qubit) and size "
            + "distribution measurements (Fragment Analyzer) "
            + "can be found as attachment to each request in "
            + "Parkour (parkour.ie-freiburg.mpg.de). Given "
            + "Library Preparation Methods are detailed in the "
            + "appendix."
        )
        lib_prep_objects = LibraryPreparation.objects.filter(
            archived=False, sample__in=instance.samples.all()
        )
        header = [
            "Date",
            "ID",
            "Name",
            "Protocol",
            "Index I7",
            "Index I5",
            "PCR",
            "ng/µl",
            "bp",
            "nM",
            "Comments",
        ]
        data = []
        for r in lib_prep_objects:
            row = [
                r.create_time.strftime("%d.%m.%Y"),
                r.sample.barcode,
                r.sample.name,
                r.sample.library_protocol.name,
                r.sample.index_i7,
                r.sample.index_i5,
                r.pcr_cycles,
                r.concentration_library,
                r.mean_fragment_size,
                r.nM,
                r.comments,
            ]
            data.append(row)
        add_table(doc, header, data)
        doc.add_page_break()

        # Page 5
        doc.add_heading("Cluster Generation and Sequencing", 1)
        doc.add_paragraph()
        header = [
            "Date",
            "ID",
            "Name",
            "Pool ID",
            "Flowcell ID",
            "Sequencer",
            "Depth (M)",
            "% Confident off species reads",
        ]
        data = []
        try:
            flowcell = instance.flowcell.get()
        except Exception:
            flowcell = None
        if flowcell:
            pool_ids = ", ".join(
                sorted(set(flowcell.lanes.values_list("pool__name", flat=True)))
            )
            sequences = flowcell.sequences if flowcell.sequences else []
            conf_reads = {s["barcode"]: s.get("confident_reads", "") for s in sequences}
            for r in records:
                row = [
                    flowcell.create_time.strftime("%d.%m.%Y"),
                    r.barcode,
                    r.name,
                    pool_ids,
                    flowcell.flowcell_id,
                    flowcell.sequencer.name,
                    r.sequencing_depth,
                    conf_reads.get(r.barcode, ""),
                ]
                data.append(row)
        add_table(doc, header, data, contains_comments=False)
        doc.add_page_break()

        # Page 6
        doc.add_heading("Acknowledgements", 1)
        doc.add_paragraph()
        doc.add_paragraph(
            "If data produced in the Deep Sequencing Facility "
            + "at MPI-IE, Freiburg is published, include an "
            + "acknowledgement in your paper. Also, review if "
            + "contributions are substantial and should lead to "
            + "an authorship of staff of the facility. "
        )
        doc.add_paragraph()
        doc.add_paragraph(
            "Additionally, let us know of any publications "
            + "involving the facility. Tracking citations and "
            + "publications demonstrate the usefulness of the "
            + "facility as a research resource which is needed "
            + "to obtain further funding."
        )
        doc.add_paragraph()
        doc.add_paragraph("Example acknowledgement")
        doc.add_paragraph()
        doc.add_paragraph(
            "We thank the Deep Sequencing Facility @ MPI-IE "
            + "Freiburg, for performance of quality controls, "
            + "library construction and Illumina sequencing."
        )
        doc.add_page_break()

        # Page 7
        doc.add_heading("Appendix", 1)
        doc.add_paragraph()
        doc.add_paragraph(
            "Detailed list of different library preparation "
            + "protocols, sequencing devices and installed "
            + "software"
        )

        doc.save(response)
        return response

    @action(methods=["get"], detail=True)
    def get_filepaths(self, request, *args, **kwargs):
        filepaths = self.get_object().filepaths
        return JsonResponse({"success": True, "filepaths": filepaths})

    @action(methods=["post"], detail=True, permission_classes=[IsAdminUser])
    def put_filepaths(self, request, pk=None):
        instance = self.get_object()
        instance.filepaths = request.data
        records = list(instance.libraries.all()) + list(instance.samples.all())
        for r in records:
            # 'Sequencing' -> 'Delivered'
            if r.status == 5:
                r.status += 1
                r.save()
        instance.save(update_fields=["filepaths"])
        return Response({"success": True})

    @action(methods=["post"], detail=True)
    def put_metapaths(self, request, pk=None):
        instance = self.get_object()
        instance.metapaths = request.data
        instance.save(update_fields=["metapaths"])
        return Response({"success": True})

    @action(methods=["get"], detail=True, permission_classes=[IsAdminUser])
    def get_flowcell(self, request, *args, **kwargs):
        instance = self.get_object()

        def get_flowcell_from_record(record, instance=instance):
            finalized = len(instance.statuses) == sum(
                [x >= 5 or x < 0 for x in instance.statuses]
            )
            if not finalized:
                value = ["Sequencing incomplete"]
            else:
                fcids = []
                for pool in record.pool.all():
                    for lane in pool.lane_set.all():
                        for flowcell in lane.flowcell.all():
                            fcids.append(flowcell)
                value = [f.flowcell_id for f in fcids]
            return json.dumps(value)

        records = list(instance.libraries.all()) + list(instance.samples.all())
        flowpaths = dict.fromkeys([r.barcode for r in records])
        for r in records:
            flowpaths[r.barcode] = get_flowcell_from_record(r)
        return JsonResponse({"flowpaths": flowpaths})

    @action(methods=["get"], detail=True, permission_classes=[IsAdminUser])
    def get_poolpaths(self, request, *args, **kwargs):
        instance = self.get_object()
        records = list(instance.libraries.all()) + list(instance.samples.all())
        poolpaths = dict.fromkeys([r.barcode for r in records])
        for r in records:
            poolpaths[r.barcode] = [p.name for p in list(r.pool.all())]
        return JsonResponse({"success": True, "poolpaths": poolpaths})

    def _get_post_data(self, request):
        post_data = {}
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            post_data = request.data.get("data", {})
            if isinstance(post_data, str):
                post_data = json.loads(post_data)
        else:
            post_data = json.loads(request.data.get("data", "{}"))
        return post_data


@login_required
def export_request(request):
    if request.method == "POST":
        primary_key = request.POST["project-id"]
        file_format = request.POST["file-format"]
        req = get_object_or_404(Request, id=primary_key)
        if not request.user.is_staff and req.user != request.user:
            raise PermissionDenied()
        dataset = Dataset()
        dataset.headers = (
            "id",
            "name",
            "barcode",
            "nucleic_acid_type",  # samples
            "library_protocol",
            "library_type",
            "concentration",
            "mean_fragment_size",  # libraries
            "index_type",  # libraries
            "index_reads",  # libraries
            "index_i7",  # libraries
            "index_i5",  # libraries
            "rna_quality",  # samples
            "read_length",
            "sequencing_depth",
            "organism",
            "comments",
        )
        records = req.records
        for r in records:
            r_type = r.__class__.__name__
            if r_type == "Sample":
                dataset.append(
                    (
                        "_",  # id
                        r.name,
                        "_",  # barcode
                        r.nucleic_acid_type,
                        r.library_protocol,
                        r.library_type,
                        r.concentration,
                        "_",  # mean_fragment_size
                        "_",  # index_type
                        "_",  # index_reads
                        "_",  # index_i7
                        "_",  # index_i5
                        r.rna_quality,
                        r.read_length,
                        r.sequencing_depth,
                        r.organism,
                        r.comments,
                    )
                )
            elif r_type == "Library":
                dataset.append(
                    (
                        "_",  # id
                        r.name,
                        "_",  # barcode
                        "_",  # nucleic_acid_type
                        r.library_protocol,
                        r.library_type,
                        r.concentration,
                        r.mean_fragment_size,
                        r.index_type,
                        r.index_reads,
                        r.index_i7,
                        r.index_i5,
                        "_",  # rna_quality
                        r.read_length,
                        r.sequencing_depth,
                        r.organism,
                        r.comments,
                    )
                )
            else:
                raise RuntimeError(f"What's {r.barcode} with {r_type}?!")

        if file_format == "CSV":
            response = HttpResponse(dataset.csv, content_type="text/csv")
            response["Content-Disposition"] = (
                'attachment; filename="exported_project_' + str(primary_key) + '.csv"'
            )
            return response

        elif file_format == "JSON":
            return render(
                request, "export.html", {"errors": "JSON import not implemented yet"}
            )

        else:
            raise ValueError(f"Invalid file_format: {file_format}")

    return render(request, "export.html")


# TODO: what about other fields from Libraries or Samples? like status, BarcodeCounter, ...
@login_required
def import_request(request):
    if request.method == "POST":
        file_format = request.POST["file-format"]
        new_file = request.FILES["importData"]

        if new_file.size > 5 * 1024 * 1024:  # 5 MB limit
            return render(
                request, "import.html", {"errors": "File size exceeds 5 MB limit"}
            )

        if file_format == "CSV":
            try:
                with transaction.atomic():
                    new_request = Request.objects.create(user=request.user)
                    csv_file = StringIO(new_file.read().decode("utf-8"))
                    csv_reader = csv.DictReader(csv_file)

                    samples = []
                    libraries = []

                    for row in csv_reader:
                        record_type = row.get("record_type", "").upper()

                        if record_type == "S":
                            sample = Sample.objects.create(
                                name=row.get("name"),
                                # Add other Sample-specific fields here
                            )
                            samples.append(sample)
                        elif record_type == "L":
                            library = Library.objects.create(
                                name=row.get("name"),
                                # Add other Library-specific fields here
                            )
                            libraries.append(library)
                        else:
                            raise ValueError(f"Invalid record_type: {record_type}")

                    # Add the samples and libraries to the request
                    new_request.samples.add(*samples)
                    new_request.libraries.add(*libraries)

                return render(request, "import.html", {"success": True})

            except Exception as e:
                return render(
                    request,
                    "import.html",
                    {"errors": f"Error processing CSV: {str(e)}"},
                )

        elif file_format == "JSON":
            return render(
                request, "import.html", {"errors": "JSON import not implemented yet"}
            )

    return render(request, "import.html")


class ApproveViewSet(viewsets.ModelViewSet):
    permission_classes = [AllowAny]
    serializer_class = RequestSerializer
    queryset = Request.objects.all().filter(id=0)

    @action(methods=["get"], detail=False)
    def this(self, request, *args, **kwargs):  # pragma: no cover
        """Process token sent to PI."""
        error = ""
        try:
            token = request.query_params.get("token")
            pk = request.query_params.get("pk")
            instance = list(Request.objects.filter(id=pk))[0]
            if not all(s == 0 for s in instance.statuses):
                raise ValueError(f"Not all statuses are zero: {instance.statuses}")
            if token == instance.token:
                instance.libraries.all().update(status=1)
                instance.samples.all().update(status=1)
                instance.token = None
                instance.approval = {
                    "TIMESTAMP": dateformat.format(timezone.now(), "c"),
                    "TOKEN": token,
                    "REMOTE_ADDR": request.META["REMOTE_ADDR"],
                    "REMOTE_PORT": request.META["REMOTE_PORT"],
                    "HTTP_USER_AGENT": request.headers["user-agent"],
                    "HTTP_ACCEPT": request.headers["accept"],
                    "HTTP_ACCEPT_ENCODING": request.headers["accept-encoding"],
                    "HTTP_ACCEPT_LANGUAGE": request.headers["accept-language"],
                    "HTTP_X_FORWARDED_FOR": request.headers["x-forwarded-for"],
                    "HTTP_X_REAL_IP": request.headers["x-real-ip"],
                }
                instance.save(update_fields=["token", "approval"])
            else:
                raise ValueError(f"Token mismatch.")
        except Exception as e:
            error = str(e)
            logger.exception(e)
            return JsonResponse({"success": not error, "error": error})
        send_mail(
            subject=f"[ Parkour2 | seq. request was approved ] {instance.name}",
            message="",
            html_message=render_to_string(
                "approved.html",
                {
                    "full_name": instance.user.full_name,
                    "pi_name": instance.user.pi.name,
                },
            ),
            from_email=settings.SERVER_EMAIL,
            recipient_list=[instance.user.email, instance.user.pi.email],
        )
        return HttpResponseRedirect("/danke")
